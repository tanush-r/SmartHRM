from fastapi import FastAPI, HTTPException
from fastapi.responses import JSONResponse, RedirectResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from gptea import QuestionAnswerChain
import pymupdf
from docx import Document
from io import BytesIO
import mysql.connector
from pydantic import BaseModel
from typing import List, Optional
from datetime import datetime
import os
import re
from dotenv import load_dotenv
import boto3
import tempfile

# Load environment variables from a .env file
load_dotenv()
app = FastAPI(root_path="/api/viewer")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
AWS_ACCESS_KEY = os.getenv('AWS_ACCESS_KEY_ID')
AWS_SECRET_KEY = os.getenv('AWS_SECRET_ACCESS_KEY')
AWS_REGION = os.getenv('AWS_REGION_NAME')
AWS_S3_BUCKET_NAME = os.getenv('AWS_S3_BUCKET_NAME')
s3 = boto3.client('s3',
                  aws_access_key_id=AWS_ACCESS_KEY,
                  aws_secret_access_key=AWS_SECRET_KEY,
                  region_name=AWS_REGION)
# Database connection
def get_db_connection():
    return mysql.connector.connect(
        host=os.getenv('MYSQL_HOST'),
        user=os.getenv('MYSQL_USER'),
        password=os.getenv('MYSQL_PASSWORD'),
        database=os.getenv('MYSQL_DB'),
        port=os.getenv('MYSQL_PORT', 3306)  # Default to 3306 if not set
    )
# Pydantic models
class Client(BaseModel):
    cl_id: str
    cl_name: str
class JobDescription(BaseModel):
    jd_id: str
    cl_id: str
    created_at: str
    s3_link: str
    filename:str
    created_at: Optional[str] = None
class Resume(BaseModel):
    resume_id: str
    jd_id: str
    s3_link: str
    filename: str
    created_at: Optional[str] = None
    st_id: Optional[str] = None  # Changed from 'status' to 'st_id'
    st_name: Optional[str] = None

class Status(BaseModel):
    st_id: str
    st_name: str

# Convert binary to hex string
def binary_to_hex(data: bytes) -> str:
    return data.hex()
# Convert datetime to string
def datetime_to_str(dt: datetime) -> str:
    return dt.isoformat()
# Get clients
@app.get("/clients", response_model=List[Client])
def get_clients():
    cursor = None
    try:
        db = get_db_connection()
        cursor = db.cursor(dictionary=True)
        cursor.execute("SELECT cl_id, cl_name FROM clients")
        clients = cursor.fetchall()
        # for client in clients:
        #     client['cl_id'] = binary_to_hex(client['cl_id'])
        return clients
    except mysql.connector.Error as err:
        raise HTTPException(status_code=500, detail=str(err))
    finally:
        if cursor:
            cursor.close()
            db.close()
# Get positions by client ID
@app.get("/positions", response_model=List[JobDescription])
def get_positions(cl_id: str):
    cursor = None
    try:
        db = get_db_connection()
        cursor = db.cursor(dictionary=True)
        cursor.execute("SELECT jd_id, cl_id, filename, s3_link, created_at FROM job_descriptions WHERE cl_id = %s", (cl_id,))
        job_descriptions = cursor.fetchall()
        for jd in job_descriptions:
            jd['created_at'] = str(jd['created_at'])
        return job_descriptions
    except mysql.connector.Error as err:
        raise HTTPException(status_code=500, detail=str(err))
    finally:
        if cursor:
            cursor.close()
            db.close()
# Get resumes by job description ID
@app.get("/resumes", response_model=List[Resume])
def get_resumes(jd_id: str):
    cursor = None
    try:
        db = get_db_connection()
        cursor = db.cursor(dictionary=True)
        cursor.execute("""
            SELECT r.resume_id, r.jd_id, r.filename, r.s3_link, r.created_at, r.st_id, s.st_name
            FROM resumes r
            LEFT JOIN status s ON r.st_id = s.st_id
            WHERE r.jd_id = %s
            ORDER BY r.created_at DESC
        """, (jd_id,))
        resumes = cursor.fetchall()

        for resume in resumes:
            if resume['created_at']:
                resume['created_at'] = datetime_to_str(resume['created_at'])
            
        print(resumes)
        
        return resumes
    except mysql.connector.Error as err:
        raise HTTPException(status_code=500, detail=str(err))
    finally:
        if cursor:
            cursor.close()
            db.close()
# Download resume by resume ID
@app.get("/resumes/download")
def download_resume(resume_id: str):
    cursor = None
    try:
        db = get_db_connection()
        cursor = db.cursor(dictionary=True)
        cursor.execute("SELECT filename FROM resumes WHERE resume_id = %s", (resume_id,))
        resumes = cursor.fetchone()
        if resumes is None:
            raise HTTPException(status_code=404, detail="Resume not found")
        filename = resumes['filename']
        # Redirect to the S3 link for downloading
        object_name = f'resumes/{filename}'
        download_path = os.path.join(tempfile.gettempdir(), filename)
        s3.download_file(AWS_S3_BUCKET_NAME, object_name, download_path)
        return FileResponse(download_path, media_type='application/octet-stream', filename=filename)
    except mysql.connector.Error as err:
        raise HTTPException(status_code=500, detail=str(err))
    finally:
        if cursor:
            cursor.close()
            db.close()
# Download resume by resume ID
@app.get("/positions/download")
def download_resume(jd_id: str):
    cursor = None
    try:
        db = get_db_connection()
        cursor = db.cursor(dictionary=True)
        cursor.execute("SELECT filename FROM job_descriptions WHERE jd_id = %s", (jd_id,))
        jds = cursor.fetchone()
        if jds is None:
            raise HTTPException(status_code=404, detail="JD not found")
        filename = jds['filename']
        # Redirect to the S3 link for downloading
        object_name = f'JD/{filename}'
        download_path = os.path.join(tempfile.gettempdir(), filename)
        s3.download_file(AWS_S3_BUCKET_NAME, object_name, download_path)
        return FileResponse(download_path, media_type='application/octet-stream', filename=filename)
    except mysql.connector.Error as err:
        raise HTTPException(status_code=500, detail=str(err))
    finally:
        if cursor:
            cursor.close()
            db.close()
allowed_files = ["pdf", "docx"]
def extract_filetype(filename: str):
    return filename.rsplit('.', 1)[1].lower()
def read_pdf(content):
    doc = pymupdf.Document(stream=content)
    # Extract the text from all pages
    text_content = ""
    for page in doc:
        text_content += page.get_text()
    return text_content
def read_docx(content):
    doc = Document(BytesIO(content))
    # Extract text from the document
    return "\n".join([para.text for para in doc.paragraphs])
@app.get("/qa_gen/resume")
async def generate_questions(count: int, resume_id: str):
    cursor = None
    try:
        db = get_db_connection()
        cursor = db.cursor(dictionary=True)
        cursor.execute("SELECT filename FROM resumes WHERE resume_id = %s", (resume_id,))
        resumes = cursor.fetchone()
        if resumes is None:
            raise HTTPException(status_code=404, detail="Resume not found")
        filename = resumes['filename']
        # Redirect to the S3 link for downloading
        object_name = f'resumes/{filename}'
        download_path = os.path.join(tempfile.gettempdir(), filename)
        s3.download_file(AWS_S3_BUCKET_NAME, object_name, download_path)
    except mysql.connector.Error as err:
        raise HTTPException(status_code=500, detail=str(err))
    finally:
        if cursor:
            cursor.close()
            db.close()
    if extract_filetype(filename) == "pdf":
        pdf_doc = pymupdf.open(download_path)
        # Extract the text from all pages
        doc = ""
        for page in pdf_doc:
            doc += page.get_text()
    elif extract_filetype(filename) == "docx":
        docx_doc = Document(download_path)
        # Extract text from the document
        doc = "\n".join([para.text for para in docx_doc.paragraphs])
    else:
        raise HTTPException(status_code=400, detail="File type not allowed. Only PDF and Word documents are accepted.")
    chain = QuestionAnswerChain(doc=doc, is_resume=True)
    return chain.invoke(count)
@app.get("/qa_gen/jd")
async def generate_questions(count: int, jd_id: str):
    cursor = None
    try:
        db = get_db_connection()
        cursor = db.cursor(dictionary=True)
        cursor.execute("SELECT filename FROM job_descriptions WHERE jd_id = %s", (jd_id,))
        jds = cursor.fetchone()
        if jds is None:
            raise HTTPException(status_code=404, detail="JD not found")
        filename = jds['filename']
        # Redirect to the S3 link for downloading
        object_name = f'JD/{filename}'
        download_path = os.path.join(tempfile.gettempdir(), filename)
        s3.download_file(AWS_S3_BUCKET_NAME, object_name, download_path)
    except mysql.connector.Error as err:
        raise HTTPException(status_code=500, detail=str(err))
    finally:
        if cursor:
            cursor.close()
            db.close()
    if extract_filetype(filename) == "pdf":
        pdf_doc = pymupdf.open(download_path)
        # Extract the text from all pages
        doc = ""
        for page in pdf_doc:
            doc += page.get_text()
    elif extract_filetype(filename) == "docx":
        docx_doc = Document(download_path)
        # Extract text from the document
        doc = "\n".join([para.text for para in docx_doc.paragraphs])
    else:
        raise HTTPException(status_code=400, detail="File type not allowed. Only PDF and Word documents are accepted.")
    chain = QuestionAnswerChain(doc=doc, is_resume=False)
    return chain.invoke(count)


# Define regex patterns for extracting name, email, and contact number
patterns = {
    'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
    'phone': r'\b(?:\+?\d{1,3}[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}\b'
}

# Extract name, email, and phone number using regex
def extract_basic_details(text):
    details = {}

    # Extract email and phone using regex
    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        details[key] = match.group(0) if match else 'N/A'
    
    # Extract name (assuming the name is the first non-empty line)
    lines = text.splitlines()
    details['name'] = next((line.strip() for line in lines if line.strip()), "Name not found")

    return details
       
@app.get("/primary")
async def primary_details(resume_id: str):

    cursor = None
    try:
        db = get_db_connection()
        cursor = db.cursor(dictionary=True)
        cursor.execute("SELECT filename FROM resumes WHERE resume_id = %s", (resume_id,))
        resumes = cursor.fetchone()
        if resumes is None:
            raise HTTPException(status_code=404, detail="Resume not found")
        
        filename = resumes['filename']
        
        # Redirect to the S3 link for downloading
        object_name = f'resumes/{filename}'

        download_path = os.path.join(tempfile.gettempdir(), filename)
        s3.download_file(AWS_S3_BUCKET_NAME, object_name, download_path)  
    except mysql.connector.Error as err:
        
        raise HTTPException(status_code=500, detail=str(err))
    finally:
        if cursor:
            cursor.close()
            db.close()

    if extract_filetype(filename) == "pdf":
        pdf_doc = pymupdf.open(download_path)

        # Extract the text from all pages
        doc = ""
        for page in pdf_doc:
            doc += page.get_text()
    elif extract_filetype(filename) == "docx":
        docx_doc = Document(download_path)

        # Extract text from the document
        doc = "\n".join([para.text for para in docx_doc.paragraphs])
    else:
        raise HTTPException(status_code=400, detail="File type not allowed. Only PDF and Word documents are accepted.")
    return extract_basic_details(doc)


# Get all statuses
@app.get("/statuses", response_model=List[Status])
async def get_all_statuses():
    """
    Fetch all available statuses from the status table.
    """
    cursor = None
    try:
        db = get_db_connection()
        cursor = db.cursor(dictionary=True)
        cursor.execute("SELECT st_id, st_name FROM status")
        statuses = cursor.fetchall()

        # If no statuses are found
        if not statuses:
            raise HTTPException(status_code=404, detail="No statuses found")

        return statuses
    except mysql.connector.Error as err:
        raise HTTPException(status_code=500, detail=str(err))
    finally:
        if cursor:
            cursor.close()
            db.close()


# Update resume status (st_id)
@app.post("/resumes/status/{resume_id}")
async def update_resume_status(resume_id: str, st_id: str):
    """
    Update the status of a resume by setting the st_id (foreign key) in the resumes table.
    """
    cursor = None
    try:
        db = get_db_connection()
        cursor = db.cursor()

        # Check if the status id exists
        cursor.execute("SELECT st_id FROM status WHERE st_id = %s", (st_id,))
        status = cursor.fetchone()
        if not status:
            raise HTTPException(status_code=404, detail="Status not found")

        # Update the st_id in the resumes table
        cursor.execute("UPDATE resumes SET st_id = %s WHERE resume_id = %s", (st_id, resume_id))
        db.commit()

        # Check if any row was updated
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="Resume not found")

        return {"message": "Resume status updated successfully"}
    except mysql.connector.Error as err:
        raise HTTPException(status_code=500, detail=str(err))
    finally:
        if cursor:
            cursor.close()
            db.close()

# Get resume status by resume ID
@app.get("/resumes/status/{resume_id}", response_model=Resume)
async def get_resume_status(resume_id: str):
    """
    Get the status of a specific resume by resume_id using st_id.
    """
    cursor = None
    try:
        db = get_db_connection()
        cursor = db.cursor(dictionary=True)
        cursor.execute("""
            SELECT r.resume_id, r.jd_id, r.filename, r.s3_link, r.created_at, r.st_id, s.st_name
            FROM resumes r
            LEFT JOIN status s ON r.st_id = s.st_id
            WHERE r.resume_id = %s
        """, (resume_id,))
        resume = cursor.fetchone()

        if not resume:
            raise HTTPException(status_code=404, detail="Resume not found")

        # Convert timestamp to string if it exists
        if resume['created_at']:
            resume['created_at'] = resume['created_at'].isoformat()

        return resume
    except mysql.connector.Error as err:
        raise HTTPException(status_code=500, detail=str(err))
    finally:
        if cursor:
            cursor.close()
            db.close()



if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)



     
